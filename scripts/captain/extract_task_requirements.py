# -*- coding: utf-8 -*-
"""
scripts/captain/extract_task_requirements.py —— 从 V3.0 队员手册 DOCX 忠实提取
T01—T09 的任务要求，生成机器可读的 docs/governance/task-requirements/T0X.yaml
与 source-manifest.json。

设计约束（来自队长治理指令，违反即视为提取无效）：

1. 只以传入的 DOCX 为权威来源；脚本内不得硬编码任何任务要求文本。
2. 不得硬编码本机绝对路径；DOCX 路径必须由命令行传入。
3. 不得把 DOCX 复制进仓库；脚本只读取，不写回源文件。
4. 同一源文件重复执行，九个 YAML 的内容必须逐字节稳定（不含时间戳）。
5. 文档未规定的字段写字符串 "not_specified"，不得凭常识补写。
6. 不得为了提取而修改项目运行依赖；PyYAML 已是仓库既有依赖，
   python-docx 缺失时自动退回 zipfile + ElementTree 解析。

术语与 Wave 语义完全来自手册第六部分「14 天共同里程碑」表：
Wave A 07/27-07/29、Wave B 07/30-08/05、Wave C 08/06-08/08、Code Freeze 08/09。

用法：

    py -3 scripts/captain/extract_task_requirements.py --docx <path-to-v3.docx>
    py -3 scripts/captain/extract_task_requirements.py --docx <path> --check

接口文档见 docs/governance/CURSOR_PR_REVIEW_RUNBOOK.md 的
「任务内容与 Wave 验收」章节。
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple
from xml.etree import ElementTree

try:  # PyYAML 是仓库既有依赖；缺失时直接失败而不是静默降级。
    import yaml
except ImportError as exc:  # pragma: no cover - 依赖缺失属环境问题
    raise SystemExit("PyYAML is required: python -m pip install -r requirements.txt") from exc

# ---------------------------------------------------------------------------
# 常量：全部为「结构定位锚点」，不是任务要求内容本身。
# ---------------------------------------------------------------------------

# 仓库常量（治理文件所属仓库，与 pr-review-policy.yaml 保持一致）。
REPOSITORY = "sage125-ai-scientist-team/SAGE125-AI-Scientist"

# 文档未规定字段的统一占位值。
NOT_SPECIFIED = "not_specified"

# 九项任务的编号集合（手册第六部分 owner map 固定为 T01—T09）。
TASK_IDS: Tuple[str, ...] = tuple("T0%d" % n for n in range(1, 10))

# 表头指纹：用于在不依赖表格序号的前提下定位四类表格。
HEADER_WAVE_MILESTONE = ("阶段", "日期", "共同目标", "完成条件")
HEADER_OWNER_MAP = ("编号", "任务", "唯一 owner 路径", "依赖", "配对审查")
HEADER_TASK_INFO = ("项目", "内容")
HEADER_DAILY_PLAN = ("日期", "重点", "具体工作", "当日交付", "验收/证据")

# 必须交付表首格的正则（例如 "T01.1 必须交付"）。
RE_MUST_DELIVER_CELL = re.compile(r"^T(0[1-9])\.1\s*必须交付")

# 任务信息表的行标签（手册对九项任务使用完全一致的九行结构）。
INFO_ROW_MISSION = "任务使命"
INFO_ROW_SCORING = "评分映射"
INFO_ROW_OWNER_PATHS = "主责路径"
INFO_ROW_FORBIDDEN = "禁止越界"
INFO_ROW_DEPENDENCIES = "依赖关系"
INFO_ROW_BRANCH_SERIES = "分支系列"
INFO_ROW_PAIRED_REVIEWER = "配对审查"
INFO_ROW_WORKLOAD = "统一工作量"

# 手册使用 □ 作为清单项前缀；分条时一并去除。
BULLET_PREFIXES = ("□", "☐", "■", "●", "•", "-", "*")

# 科学真实性 / 安全类关键词：命中即升级为 P0_BLOCKING（造假与泄密零容忍）。
P0_PATTERNS: Tuple[str, ...] = (
    "不得把模拟指标写成真实结果",
    "不得将 Mock 标为正式结果",
    "Mock 标为正式结果",
    "不得虚构",
    "虚构引用",
    "不得伪造",
    "不得为 Mock",
    "不得为 Mock",
    "正式结果不得为 Mock",
    "不以文件名或 OCR 文本冒充多模态",
    "不得在 UI 层伪造状态或结果",
    "不得报告为成功",
    "不得暴露密钥",
    "不暴露密钥",
    "密钥不进镜像",
    "无 API key",
    "不写入真实密钥",
    "actual_execution",
    "不得通过删除断言",
    "不得降低测试标准",
)

# 明确选做的标记；只有文档自己这样写才允许 OPTIONAL。
OPTIONAL_PATTERNS: Tuple[str, ...] = ("可选", "选做", "如适用", "可跳过")

# 硬要求动词：出现任一即不得降级为 P2 建议。
HARD_REQUIREMENT_MARKERS: Tuple[str, ...] = (
    "必须", "不得", "禁止", "至少", "不低于", "不少于", "不高于", "不超过",
    "应", "须", "完成", "实现", "交付", "建立", "运行", "提交", "确认",
)

# 纯建议标记：仅当同时不含硬要求动词时才判为 P2_RECOMMENDATION。
RECOMMENDATION_ONLY_MARKERS: Tuple[str, ...] = ("建议", "推荐", "可进一步", "可增强", "宜")

# verification_type 关键词映射（顺序敏感：先命中者胜）。
VERIFICATION_RULES: Tuple[Tuple[str, Tuple[str, ...]], ...] = (
    (
        "metric_reproduction",
        ("精确率", "recall", "准确率", "误差", "相对误差", "指标", "metrics", "召回",
         "通过率", "延迟", "token", "成本", "覆盖率", "相似度", "AUROC", "评测", "消融"),
    ),
    (
        "test_execution",
        ("测试", "红灯", "夹具", "fixture", "E2E", "回归", "用例", "CI", "job",
         "workflow", "扫描", "校验器", "validator", "抽查"),
    ),
    (
        "artifact_inspection",
        ("交付", "产物", "报告", "manifest", "checksum", "SHA-256", "trace", "截图",
         "导出", "release", "handoff", "目录", "索引", "日志", "图表", "PDF", "JSON"),
    ),
    (
        "doc_inspection",
        ("文档", "说明", "docs/", "README", "RFC", "记录", "笔记", "清单", "模板", "迁移"),
    ),
    (
        "manual_signoff",
        ("人工核验", "签字", "人工审查", "人工标注", "人工/规则校验"),
    ),
)

# evidence_required 关键词映射：命中即追加对应证据类别。
EVIDENCE_RULES: Tuple[Tuple[str, Tuple[str, ...]], ...] = (
    ("test_output", ("测试", "红灯", "用例", "回归", "E2E", "CI", "job", "扫描", "抽查")),
    ("metrics_file", ("精确率", "recall", "准确率", "误差", "指标", "metrics", "通过率",
                      "延迟", "成本", "覆盖率", "相似度")),
    ("computation_script", ("脚本", "聚合", "计算", "评测", "消融", "benchmark", "校验器")),
    ("raw_results", ("原始", "raw", "JSON", "CSV", "日志", "stdout", "stderr", "trace")),
    ("dataset_manifest", ("黄金集", "数据集", "题集", "样本", "manifest", "许可", "数据来源")),
    ("artifact_file", ("产物", "报告", "导出", "PDF", "MD", "release", "checksum", "SHA-256",
                       "图表", "截图", "目录", "索引")),
    ("source_code", ("实现", "接口", "契约", "Schema", "字段", "序列化", "运行器", "API",
                     "解析", "检查器", "渲染", "路由", "前端", "workspace")),
    ("documentation", ("文档", "说明", "docs/", "README", "RFC", "记录", "笔记", "迁移", "回滚")),
    ("reproduction_command", ("复现", "可复跑", "可再生成", "干净环境", "clone", "运行命令")),
    ("git_commit_reference", ("commit", "Git", "分支", "PR")),
)

# 手册中出现的中文数词；仅用于「至少一个」这类没有阿拉伯数字的阈值。
CHINESE_NUMERALS: Dict[str, int] = {
    "一": 1, "两": 2, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}

# 计量单位白名单：只有紧跟这些单位的裸数字才视为定量阈值，避免把日期和编号误判。
QUANTITY_UNITS = "条|个|题|种|类|页|次|小时|并发|领域|轮|查询|文件|基线|问题|样本|维|项|天"

# 定量阈值识别规则：(comparator, 正则)。按顺序匹配，已消费的字符区间不再复用。
# 只识别文档写明的数值，不做任何推断、换算或补全。
THRESHOLD_RULES: Tuple[Tuple[str, "re.Pattern[str]"], ...] = (
    ("ratio", re.compile(r"\b([0-9]+)\s*/\s*([0-9]+)\b")),
    (
        "at_least",
        re.compile(
            r"(?:不低于|不少于|至少|大于等于|≥|>=)\s*([0-9]+(?:\.[0-9]+)?)\s*([%％]|" + QUANTITY_UNITS + r")?"
        ),
    ),
    (
        "at_most",
        re.compile(
            r"(?:不超过|不高于|小于等于|≤|<=)\s*([0-9]+(?:\.[0-9]+)?)\s*([%％]|" + QUANTITY_UNITS + r"|秒)?"
        ),
    ),
    ("greater_than", re.compile(r"(?:大于|>)\s*([0-9]+(?:\.[0-9]+)?)\s*([%％])?")),
    ("exact", re.compile(r"(?:total\s*=\s*|共\s*)([0-9]+)\s*([%％])?")),
    ("percentage", re.compile(r"([0-9]+(?:\.[0-9]+)?)\s*([%％])")),
    (
        "at_least",
        re.compile(
            r"(?:至少|不少于)\s*([" + "".join(CHINESE_NUMERALS) + r"])\s*(" + QUANTITY_UNITS + r")"
        ),
    ),
    ("quantity", re.compile(r"([0-9]+(?:\.[0-9]+)?)\s*(" + QUANTITY_UNITS + r")")),
)


# ---------------------------------------------------------------------------
# DOCX 读取层
# ---------------------------------------------------------------------------


class DocxContent:
    """
    DOCX 解析结果的只读容器。

    属性：
        paragraphs: [(index, style_name, text)]，仅保留非空段落但保留原始索引。
        tables:     [[[cell_text, ...], ...]]，按文档顺序排列的表格。
    """

    def __init__(
        self,
        paragraphs: List[Tuple[int, str, str]],
        tables: List[List[List[str]]],
    ) -> None:
        """保存段落与表格；本类不做任何语义解释。"""
        self.paragraphs = paragraphs
        self.tables = tables


def _read_docx_with_python_docx(path: Path) -> Optional[DocxContent]:
    """
    使用 python-docx 读取 DOCX。

    返回：
        解析成功返回 DocxContent；python-docx 不可用时返回 None，
        由调用方回退到 zipfile 解析（不修改项目依赖）。
    """
    try:
        import docx  # type: ignore
    except ImportError:
        return None
    document = docx.Document(str(path))
    paragraphs = [
        (index, paragraph.style.name if paragraph.style is not None else "", paragraph.text.strip())
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]
    tables = [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    return DocxContent(paragraphs, tables)


def _read_docx_with_zipfile(path: Path) -> DocxContent:
    """
    不依赖 python-docx 的 DOCX 解析回退实现。

    直接读取 word/document.xml，按 WordprocessingML 语义还原段落与表格：
    w:p 为段落、w:tbl 为表格、w:tr 为行、w:tc 为单元格、w:t 为文本，
    w:br / w:cr 视为换行以保留手册中的 □ 分条结构。
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(str(path)) as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml_bytes)
    body = root.find("w:body", ns)
    if body is None:
        raise ValueError("DOCX 缺少 w:body，文件可能已损坏")

    def paragraph_text(node: ElementTree.Element) -> str:
        """把一个 w:p 节点还原为纯文本，保留显式换行。"""
        chunks: List[str] = []
        for child in node.iter():
            tag = child.tag.split("}")[-1]
            if tag == "t" and child.text:
                chunks.append(child.text)
            elif tag in ("br", "cr"):
                chunks.append("\n")
            elif tag == "tab":
                chunks.append("\t")
        return "".join(chunks).strip()

    def paragraph_style(node: ElementTree.Element) -> str:
        """读取段落样式名；无样式返回空串。"""
        style = node.find("w:pPr/w:pStyle", ns)
        if style is None:
            return ""
        return style.attrib.get("{%s}val" % ns["w"], "")

    paragraphs: List[Tuple[int, str, str]] = []
    tables: List[List[List[str]]] = []
    paragraph_index = 0
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = paragraph_text(child)
            if text:
                paragraphs.append((paragraph_index, paragraph_style(child), text))
            paragraph_index += 1
        elif tag == "tbl":
            rows: List[List[str]] = []
            for row_node in child.findall("w:tr", ns):
                cells: List[str] = []
                for cell_node in row_node.findall("w:tc", ns):
                    cell_parts = [paragraph_text(p) for p in cell_node.findall("w:p", ns)]
                    cells.append("\n".join(part for part in cell_parts if part).strip())
                rows.append(cells)
            tables.append(rows)
    return DocxContent(paragraphs, tables)


def read_docx(path: Path) -> DocxContent:
    """
    读取 DOCX，优先 python-docx，缺失时回退 zipfile 解析。

    异常：
        FileNotFoundError: 源文件不存在（调用方必须停止，不得凭记忆编造要求）。
    """
    if not path.is_file():
        raise FileNotFoundError("V3.0 源文档不存在：%s" % path)
    content = _read_docx_with_python_docx(path)
    if content is None:
        content = _read_docx_with_zipfile(path)
    return content


# ---------------------------------------------------------------------------
# 通用文本工具
# ---------------------------------------------------------------------------


def sha256_text(text: str) -> str:
    """对 UTF-8 编码后的文本求 SHA-256，用于要求条目的稳定指纹。"""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def sha256_file(path: Path) -> str:
    """分块求文件 SHA-256，避免一次性读入大文件。"""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(65536), b""):
            digest.update(block)
    return digest.hexdigest()


def strip_bullet(line: str) -> str:
    """去掉手册清单项的 □ 等前缀与首尾空白，其余字符逐字保留。"""
    text = line.strip()
    for prefix in BULLET_PREFIXES:
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
            break
    return text


def split_bullets(cell_text: str) -> List[str]:
    """
    把单元格文本按行拆成清单项。

    手册在一个单元格内用换行 + □ 表达多条要求；空行与纯标题行被忽略。
    """
    items: List[str] = []
    for raw_line in cell_text.splitlines():
        line = strip_bullet(raw_line)
        if not line:
            continue
        items.append(line)
    return items


def split_semicolon_list(text: str) -> List[str]:
    """按中英文分号/顿号拆分枚举型单元格（如评分映射、主责路径）。"""
    parts = re.split(r"[;；]", text)
    return [part.strip().rstrip("。") for part in parts if part.strip()]


# ---------------------------------------------------------------------------
# 表格定位
# ---------------------------------------------------------------------------


def header_matches(rows: Sequence[Sequence[str]], header: Sequence[str]) -> bool:
    """判断表格首行是否与给定表头指纹一致（去空白后逐列比较）。"""
    if not rows:
        return False
    first = [cell.strip() for cell in rows[0]]
    if len(first) < len(header):
        return False
    return all(first[i] == header[i] for i in range(len(header)))


def find_table(content: DocxContent, header: Sequence[str]) -> Tuple[int, List[List[str]]]:
    """
    按表头指纹查找唯一表格。

    异常：
        LookupError: 未找到，说明文档结构与预期不符，必须停止而不是猜测。
    """
    for index, rows in enumerate(content.tables):
        if header_matches(rows, header):
            return index, rows
    raise LookupError("未找到表头为 %s 的表格" % (" | ".join(header),))


def find_must_deliver_tables(content: DocxContent) -> Dict[str, int]:
    """
    定位九张「T0X.1 必须交付」表，返回 {task_id: table_index}。

    异常：
        LookupError: 九项任务未全部定位到，提取必须中止。
    """
    found: Dict[str, int] = {}
    for index, rows in enumerate(content.tables):
        if not rows or not rows[0]:
            continue
        match = RE_MUST_DELIVER_CELL.match(rows[0][0].strip())
        if match:
            found["T" + match.group(1)] = index
    missing = [task_id for task_id in TASK_IDS if task_id not in found]
    if missing:
        raise LookupError("以下任务缺少「必须交付」表：%s" % ", ".join(missing))
    return found


def previous_table_with_header(
    content: DocxContent, before_index: int, header: Sequence[str]
) -> Tuple[int, List[List[str]]]:
    """向前查找最近一个匹配表头的表格（用于定位任务信息表）。"""
    for index in range(before_index - 1, -1, -1):
        if header_matches(content.tables[index], header):
            return index, content.tables[index]
    raise LookupError("表 #%d 之前未找到表头 %s" % (before_index, " | ".join(header)))


def next_table_with_header(
    content: DocxContent, after_index: int, header: Sequence[str]
) -> Tuple[int, List[List[str]]]:
    """向后查找最近一个匹配表头的表格（用于定位每日计划表）。"""
    for index in range(after_index + 1, len(content.tables)):
        if header_matches(content.tables[index], header):
            return index, content.tables[index]
    raise LookupError("表 #%d 之后未找到表头 %s" % (after_index, " | ".join(header)))


def find_task_heading(content: DocxContent, task_id: str) -> Tuple[int, str]:
    """
    定位任务的一级标题段落，返回 (paragraph_index, heading_text)。

    手册中标题形如 "T01. 队员 1（姓名待填）— 证据全文注入、事实蕴含与可追溯引用"。
    """
    prefix = task_id + "."
    for index, style, text in content.paragraphs:
        if text.startswith(prefix) and "Heading" in style:
            return index, text
    for index, _style, text in content.paragraphs:
        if text.startswith(prefix):
            return index, text
    raise LookupError("未找到 %s 的任务标题段落" % task_id)


# ---------------------------------------------------------------------------
# Wave 与 owner map
# ---------------------------------------------------------------------------


def parse_wave_milestones(content: DocxContent) -> Dict[str, Dict[str, str]]:
    """
    解析「14 天共同里程碑」表，得到 Wave A/B/C 与 Code Freeze 的日期窗口。

    返回：
        {"A": {"label","date_range","team_goal","completion_condition","dates":[MM/DD,...]}, ...,
         "FREEZE": {...}}
    """
    _index, rows = find_table(content, HEADER_WAVE_MILESTONE)
    waves: Dict[str, Dict[str, str]] = {}
    for row in rows[1:]:
        label = row[0].strip()
        date_range = row[1].strip()
        key = None
        if label.startswith("Wave A"):
            key = "A"
        elif label.startswith("Wave B"):
            key = "B"
        elif label.startswith("Wave C"):
            key = "C"
        elif "Code Freeze" in label:
            key = "FREEZE"
        if key is None:
            continue
        waves[key] = {
            "label": label,
            "date_range": date_range,
            "team_goal": row[2].strip(),
            "completion_condition": row[3].strip(),
            "dates": expand_date_range(date_range),
        }
    missing = [key for key in ("A", "B", "C", "FREEZE") if key not in waves]
    if missing:
        raise LookupError("里程碑表缺少阶段：%s" % ", ".join(missing))
    return waves


def expand_date_range(date_range: str) -> List[str]:
    """
    把 "07/27-07/29" 展开为 ["07/27","07/28","07/29"]；单日直接返回单元素列表。

    仅支持手册使用的 MM/DD 形式；解析失败返回空列表，交由调用方报错。
    """
    text = date_range.replace("－", "-").replace("—", "-").strip()
    single = re.fullmatch(r"(\d{2})/(\d{2})", text)
    if single:
        return [text]
    span = re.fullmatch(r"(\d{2})/(\d{2})\s*-\s*(\d{2})/(\d{2})", text)
    if not span:
        return []
    # 手册周期固定在 2026 年内，用序数逐日递增即可跨月安全展开。
    start = datetime(2026, int(span.group(1)), int(span.group(2)))
    end = datetime(2026, int(span.group(3)), int(span.group(4)))
    if end < start:
        return []
    dates: List[str] = []
    for ordinal in range(start.toordinal(), end.toordinal() + 1):
        dates.append(datetime.fromordinal(ordinal).strftime("%m/%d"))
    return dates


def parse_owner_map(content: DocxContent) -> Dict[str, Dict[str, str]]:
    """
    解析九项任务 owner map 表，返回 {task_id: {title, owner_paths, dependencies, reviewer}}。

    该表与任务信息表内容重叠，用于交叉校验提取是否一致。
    """
    _index, rows = find_table(content, HEADER_OWNER_MAP)
    owners: Dict[str, Dict[str, str]] = {}
    for row in rows[1:]:
        task_id = row[0].strip()
        if task_id not in TASK_IDS:
            continue
        owners[task_id] = {
            "title": row[1].strip(),
            "owner_paths": row[2].strip(),
            "dependencies": row[3].strip(),
            "paired_reviewer": row[4].strip(),
        }
    return owners


def parse_info_table(rows: Sequence[Sequence[str]]) -> Dict[str, str]:
    """把任务信息表（项目 | 内容）转成 {行标签: 内容} 字典。"""
    info: Dict[str, str] = {}
    for row in rows[1:]:
        if len(row) < 2:
            continue
        info[row[0].strip()] = row[1].strip()
    return info


def parse_dependencies(raw: str) -> Dict[str, Any]:
    """
    解析依赖关系文本，抽出 upstream / downstream 任务编号。

    原文保留在 raw 字段；无法识别方向时对应列表为空，不做猜测。
    """
    upstream: List[str] = []
    downstream: List[str] = []
    for segment in re.split(r"[;；]", raw):
        ids = re.findall(r"T0[1-9]", segment)
        if not ids:
            continue
        if "上游" in segment or "消费" in segment:
            upstream.extend(ids)
        elif "下游" in segment:
            downstream.extend(ids)
        else:
            upstream.extend(ids)
    dedup = lambda items: sorted(set(items))  # noqa: E731 - 保持排序稳定
    return {"raw": raw, "upstream": dedup(upstream), "downstream": dedup(downstream)}


def parse_sprint_days(workload_text: str) -> Any:
    """从统一工作量文本中提取冲刺日数；未写明返回 not_specified。"""
    match = re.search(r"(\d+)\s*个?冲刺日", workload_text)
    if match:
        return int(match.group(1))
    return NOT_SPECIFIED


def parse_branch_series(text: str) -> List[str]:
    """把 "a → b → c" 形式的分支系列拆成有序列表。"""
    parts = re.split(r"[→>]+", text)
    return [part.strip() for part in parts if part.strip()]


# ---------------------------------------------------------------------------
# 要求分类
# ---------------------------------------------------------------------------


def classify_blocking_policy(text: str, category: str) -> str:
    """
    按队长指令的统一策略为一条要求赋 blocking_policy。

    优先级：科学真实性/安全 > 文档明确选做 > 纯建议 > 类别默认（P1_BLOCKING）。
    绝不把明确的硬要求降级为 P2。
    """
    if any(pattern in text for pattern in P0_PATTERNS):
        return "P0_BLOCKING"
    if any(pattern in text for pattern in OPTIONAL_PATTERNS):
        return "OPTIONAL"
    has_hard = any(marker in text for marker in HARD_REQUIREMENT_MARKERS)
    only_recommend = any(marker in text for marker in RECOMMENDATION_ONLY_MARKERS)
    if only_recommend and not has_hard:
        return "P2_RECOMMENDATION"
    del category  # 所有内容类别的默认阻断级别一致，保留参数便于未来细化。
    return "P1_BLOCKING"


def classify_future_wave_policy(blocking_policy: str) -> str:
    """
    决定该要求在「当前 Wave 早于要求 Wave」时能否 DEFERRED。

    P0（造假与泄密）在任何 Wave 都必须评估，不可延期；其余允许延期。
    """
    if blocking_policy == "P0_BLOCKING":
        return "NOT_DEFERRABLE"
    return "DEFERRED_ALLOWED"


def classify_verification_type(text: str) -> str:
    """按关键词映射推断验证方式；无命中时保守归为 code_inspection。"""
    for verification_type, keywords in VERIFICATION_RULES:
        if any(keyword in text for keyword in keywords):
            return verification_type
    return "code_inspection"


def classify_evidence_required(text: str, verification_type: str) -> List[str]:
    """
    推断该要求必须观察到的证据类别。

    定量类要求强制包含指标复现所需的全套证据，避免「只有一个百分数」即通过。
    """
    evidence: List[str] = []
    for label, keywords in EVIDENCE_RULES:
        if any(keyword in text for keyword in keywords):
            evidence.append(label)
    if verification_type == "metric_reproduction":
        for label in (
            "dataset_manifest",
            "computation_script",
            "raw_results",
            "metrics_file",
            "reproduction_command",
            "git_commit_reference",
        ):
            if label not in evidence:
                evidence.append(label)
    if not evidence:
        evidence.append("source_code")
    return evidence


def extract_thresholds(text: str) -> List[Dict[str, Any]]:
    """
    从要求文本中提取文档写明的定量阈值。

    只识别显式出现的数值与比较词；不推断、不补全、不换算。
    返回的每一项都保留 raw_text 以便人工复核。
    """
    thresholds: List[Dict[str, Any]] = []
    consumed: List[Tuple[int, int]] = []

    def overlaps(start: int, end: int) -> bool:
        """判断当前匹配是否与更高优先级规则已占用的区间重叠。"""
        return any(start < used_end and end > used_start for used_start, used_end in consumed)

    for comparator, pattern in THRESHOLD_RULES:
        for match in pattern.finditer(text):
            if overlaps(match.start(), match.end()):
                continue
            if comparator == "ratio":
                value: Any = "%s/%s" % (match.group(1), match.group(2))
                unit = ""
            else:
                raw_value = match.group(1)
                if raw_value in CHINESE_NUMERALS:
                    value = CHINESE_NUMERALS[raw_value]
                else:
                    numeric = float(raw_value)
                    value = int(numeric) if numeric.is_integer() else numeric
                unit = (match.group(2) or "").strip() if match.re.groups >= 2 else ""
                if unit in ("%", "％"):
                    unit = "%"
            consumed.append((match.start(), match.end()))
            thresholds.append(
                {
                    "raw_text": match.group(0).strip(),
                    "comparator": comparator,
                    "value": value,
                    "unit": unit if unit else NOT_SPECIFIED,
                }
            )
    thresholds.sort(key=lambda item: text.find(item["raw_text"]))
    return thresholds


# ---------------------------------------------------------------------------
# 要求条目构建
# ---------------------------------------------------------------------------


class RequirementBuilder:
    """
    单个任务的要求条目构建器，负责稳定 ID 分配与字段填充。

    ID 形如 T01-A-001 / T01-DOD-002 / T01-METRIC-003，
    序号按文档出现顺序在各前缀内独立递增，保证重复执行结果一致。
    """

    def __init__(self, task_id: str, source: Dict[str, Any]) -> None:
        """记录任务编号与源文档信息，并初始化各前缀计数器。"""
        self.task_id = task_id
        self.source = source
        self._counters: Dict[str, int] = {}
        self.requirements: List[Dict[str, Any]] = []
        self.thresholds: List[Dict[str, Any]] = []

    def _next_id(self, prefix: str) -> str:
        """为给定前缀分配下一个三位序号。"""
        self._counters[prefix] = self._counters.get(prefix, 0) + 1
        return "%s-%s-%03d" % (self.task_id, prefix, self._counters[prefix])

    def add(
        self,
        prefix: str,
        wave: str,
        category: str,
        text: str,
        source_heading: str,
        table_index: int,
        row_index: int,
        cell_index: int,
    ) -> Dict[str, Any]:
        """
        追加一条要求，并自动派生其定量指标子要求。

        参数：
            prefix:        ID 前缀（A/B/C/MUST/DOD/HANDOFF/FREEZE）。
            wave:          该要求进入审核范围的 Wave。
            category:      要求类别（见 task-requirement.schema.json）。
            text:          手册原文（逐字，不改写）。
            source_heading: 可读的来源路径，便于人工回查文档。
            table_index/row_index/cell_index: 精确来源坐标。

        返回：
            新建的要求字典（已加入 self.requirements）。
        """
        blocking_policy = classify_blocking_policy(text, category)
        verification_type = classify_verification_type(text)
        requirement: Dict[str, Any] = {
            "id": self._next_id(prefix),
            "task_id": self.task_id,
            "wave": wave,
            "category": category,
            "requirement_text": text,
            "source_heading": source_heading,
            "source_paragraph_index": self.source["task_heading_paragraph_index"],
            "source_table_index": table_index,
            "source_row_index": row_index,
            "source_cell_index": cell_index,
            "source_text_sha256": sha256_text(text),
            "evidence_required": classify_evidence_required(text, verification_type),
            "verification_type": verification_type,
            "blocking_policy": blocking_policy,
            "future_wave_policy": classify_future_wave_policy(blocking_policy),
            "related_metric_ids": [],
        }
        self.requirements.append(requirement)
        self._add_metrics(requirement, source_heading, table_index, row_index, cell_index)
        return requirement

    def _add_metrics(
        self,
        parent: Dict[str, Any],
        source_heading: str,
        table_index: int,
        row_index: int,
        cell_index: int,
    ) -> None:
        """
        为父要求中出现的每个定量阈值生成独立的 METRIC 要求。

        定量要求一律要求完整复现证据链，因此 verification_type 固定为
        metric_reproduction，evidence_required 由 classify_evidence_required 补全。
        """
        for threshold in extract_thresholds(parent["requirement_text"]):
            metric_text = "%s（定量阈值：%s）" % (parent["requirement_text"], threshold["raw_text"])
            blocking_policy = parent["blocking_policy"]
            if blocking_policy == "OPTIONAL":
                metric_policy = "OPTIONAL"
            elif blocking_policy == "P2_RECOMMENDATION":
                metric_policy = "P2_RECOMMENDATION"
            else:
                metric_policy = "P1_BLOCKING" if blocking_policy != "P0_BLOCKING" else "P0_BLOCKING"
            metric = {
                "id": self._next_id("METRIC"),
                "task_id": self.task_id,
                "wave": parent["wave"],
                "category": "quantitative_metric",
                "requirement_text": metric_text,
                "source_heading": source_heading,
                "source_paragraph_index": self.source["task_heading_paragraph_index"],
                "source_table_index": table_index,
                "source_row_index": row_index,
                "source_cell_index": cell_index,
                "source_text_sha256": sha256_text(metric_text),
                "evidence_required": classify_evidence_required(
                    parent["requirement_text"], "metric_reproduction"
                ),
                "verification_type": "metric_reproduction",
                "blocking_policy": metric_policy,
                "future_wave_policy": classify_future_wave_policy(metric_policy),
                "related_metric_ids": [],
                "derived_from_requirement_id": parent["id"],
            }
            self.requirements.append(metric)
            parent["related_metric_ids"].append(metric["id"])
            self.thresholds.append(
                {
                    "requirement_id": metric["id"],
                    "raw_text": threshold["raw_text"],
                    "comparator": threshold["comparator"],
                    "value": threshold["value"],
                    "unit": threshold["unit"],
                    "source_requirement_id": parent["id"],
                }
            )


# ---------------------------------------------------------------------------
# 单任务提取
# ---------------------------------------------------------------------------


def wave_for_date(date_cell: str, waves: Dict[str, Dict[str, str]]) -> Optional[str]:
    """
    根据每日计划行的日期单元格判断所属 Wave。

    日期形如 "07/28（周二）"；无法匹配任何 Wave 窗口时返回 None，
    调用方必须报错而不是猜测归属。
    """
    match = re.match(r"(\d{2}/\d{2})", date_cell.strip())
    if not match:
        return None
    date = match.group(1)
    for key in ("A", "B", "C", "FREEZE"):
        if date in waves[key]["dates"]:
            return key
    return None


def build_handoff_requirements(content: DocxContent) -> List[Tuple[str, int]]:
    """
    提取第 7.3 节「8 月 9 日 handoff 目录」的全队 handoff 要求。

    返回 [(requirement_text, paragraph_index)]，九项任务共享这些条目。
    异常：
        LookupError: 未找到 7.3 节，说明文档结构异常。
    """
    start_index: Optional[int] = None
    for index, _style, text in content.paragraphs:
        if text.startswith("7.3") and "handoff" in text:
            start_index = index
            break
    if start_index is None:
        raise LookupError("未找到 7.3 handoff 章节")
    items: List[Tuple[str, int]] = []
    for index, style, text in content.paragraphs:
        if index <= start_index:
            continue
        if "Heading" in style:
            break
        cleaned = strip_bullet(text)
        if cleaned:
            items.append((cleaned, index))
    if not items:
        raise LookupError("7.3 handoff 章节没有可提取的清单项")
    return items


def extract_task(
    content: DocxContent,
    task_id: str,
    must_deliver_table_index: int,
    waves: Dict[str, Dict[str, str]],
    owner_map: Dict[str, Dict[str, str]],
    handoff_items: List[Tuple[str, int]],
    source_meta: Dict[str, Any],
) -> Dict[str, Any]:
    """
    提取单个任务的完整规范，返回可直接写入 T0X.yaml 的字典。

    提取顺序固定为：信息表 → 必须交付/DoD 表 → 每日计划表 → handoff，
    以保证 requirement ID 在重复执行时稳定。

    异常：
        LookupError: 任一必需表格或章节缺失。
        ValueError:  每日计划行的日期无法归属到任何 Wave。
    """
    heading_index, heading_text = find_task_heading(content, task_id)
    info_index, info_rows = previous_table_with_header(
        content, must_deliver_table_index, HEADER_TASK_INFO
    )
    daily_index, daily_rows = next_table_with_header(
        content, must_deliver_table_index, HEADER_DAILY_PLAN
    )
    info = parse_info_table(info_rows)
    must_rows = content.tables[must_deliver_table_index]

    source = {
        "docx_name": source_meta["docx_name"],
        "docx_sha256": source_meta["docx_sha256"],
        "docx_size_bytes": source_meta["docx_size_bytes"],
        "document_version": source_meta["document_version"],
        "execution_period": source_meta["execution_period"],
        "task_heading_paragraph_index": heading_index,
        "info_table_index": info_index,
        "must_deliver_table_index": must_deliver_table_index,
        "daily_table_index": daily_index,
    }
    builder = RequirementBuilder(task_id, source)

    # --- 必须交付（T0X.1）：整任务范围，累计在 Wave C 前全部达成。 ---
    must_deliver_items = split_bullets(must_rows[0][0])
    must_deliver_items = [
        item for item in must_deliver_items if not RE_MUST_DELIVER_CELL.match(item)
    ]
    for row_offset, item in enumerate(must_deliver_items):
        builder.add(
            prefix="MUST",
            wave="FINAL",
            category="must_deliver",
            text=item,
            source_heading="%s › %s.1 必须交付" % (heading_text, task_id),
            table_index=must_deliver_table_index,
            row_index=0,
            cell_index=0,
        )
        del row_offset

    # --- Definition of Done（T0X.2）：Wave C 逐条核验。 ---
    dod_cell = must_rows[0][1] if len(must_rows[0]) > 1 else ""
    dod_items = [
        item
        for item in split_bullets(dod_cell)
        if not re.match(r"^T0[1-9]\.2\s*Definition of Done", item)
    ]
    for item in dod_items:
        builder.add(
            prefix="DOD",
            wave="C",
            category="definition_of_done",
            text=item,
            source_heading="%s › %s.2 Definition of Done" % (heading_text, task_id),
            table_index=must_deliver_table_index,
            row_index=0,
            cell_index=1,
        )

    # --- 每日计划：具体工作 / 当日交付 / 验收证据，按 Wave 归属。 ---
    wave_days: Dict[str, List[Dict[str, str]]] = {"A": [], "B": [], "C": []}
    code_freeze: Optional[Dict[str, str]] = None
    for row_index, row in enumerate(daily_rows[1:], start=1):
        if len(row) < 5:
            continue
        date_cell, focus, work, deliverable, evidence = (cell.strip() for cell in row[:5])
        wave_key = wave_for_date(date_cell, waves)
        if wave_key is None:
            raise ValueError("%s 每日计划行日期无法归属 Wave：%s" % (task_id, date_cell))
        day_record = {
            "date": date_cell,
            "focus": focus,
            "work": work,
            "deliverable": deliverable,
            "evidence": evidence,
        }
        heading_path = "%s › %s.3 每日计划 › %s" % (heading_text, task_id, date_cell)
        if wave_key == "FREEZE":
            code_freeze = day_record
            for cell_index, (category, text) in enumerate(
                (
                    ("code_freeze_constraint", work),
                    ("code_freeze_constraint", deliverable),
                    ("code_freeze_constraint", evidence),
                ),
                start=2,
            ):
                if text:
                    builder.add(
                        prefix="FREEZE",
                        wave="FREEZE",
                        category=category,
                        text=text,
                        source_heading=heading_path,
                        table_index=daily_index,
                        row_index=row_index,
                        cell_index=cell_index,
                    )
            continue
        wave_days[wave_key].append(day_record)
        for cell_index, (category, text) in enumerate(
            (
                ("daily_work", work),
                ("daily_deliverable", deliverable),
                ("acceptance_evidence", evidence),
            ),
            start=2,
        ):
            if text:
                builder.add(
                    prefix=wave_key,
                    wave=wave_key,
                    category=category,
                    text=text,
                    source_heading=heading_path,
                    table_index=daily_index,
                    row_index=row_index,
                    cell_index=cell_index,
                )

    if code_freeze is None:
        raise LookupError("%s 每日计划缺少 Code Freeze（08/09）行" % task_id)
    for wave_key in ("A", "B", "C"):
        if not wave_days[wave_key]:
            raise LookupError("%s 缺少 Wave %s 的每日计划行" % (task_id, wave_key))

    # --- handoff：全队共享的 7.3 节要求（每个任务独立登记一份）。 ---
    for text, paragraph_index in handoff_items:
        requirement = builder.add(
            prefix="HANDOFF",
            wave="FREEZE",
            category="handoff",
            text=text,
            source_heading="第七部分 › 7.3 8 月 9 日 handoff 目录",
            table_index=0,
            row_index=0,
            cell_index=0,
        )
        requirement["source_paragraph_index"] = paragraph_index

    owner_entry = owner_map.get(task_id, {})
    forbidden_text = info.get(INFO_ROW_FORBIDDEN, "").strip()
    conflict_parts = [part for part in (forbidden_text,) if part]
    conflict_parts.append(
        "全局停止条件：冲突涉及其他任务路径、队长专属文件、破坏性公共 Schema、"
        "真实密钥或无许可数据时，必须停止自行处理并联系队长（手册 5.x）。"
    )

    document: Dict[str, Any] = {
        "schema_version": 1,
        "repository": REPOSITORY,
        "task_id": task_id,
        "title": heading_text,
        "mission": info.get(INFO_ROW_MISSION, NOT_SPECIFIED) or NOT_SPECIFIED,
        "scoring_dimensions": split_semicolon_list(info.get(INFO_ROW_SCORING, "")) or [NOT_SPECIFIED],
        "owner_paths": split_semicolon_list(info.get(INFO_ROW_OWNER_PATHS, "")) or [NOT_SPECIFIED],
        "forbidden_paths": split_semicolon_list(forbidden_text) or [NOT_SPECIFIED],
        "dependencies": parse_dependencies(
            info.get(INFO_ROW_DEPENDENCIES, owner_entry.get("dependencies", NOT_SPECIFIED))
        ),
        "branch_series": parse_branch_series(info.get(INFO_ROW_BRANCH_SERIES, "")) or [NOT_SPECIFIED],
        "paired_reviewer": info.get(
            INFO_ROW_PAIRED_REVIEWER, owner_entry.get("paired_reviewer", NOT_SPECIFIED)
        )
        or NOT_SPECIFIED,
        "sprint_days": parse_sprint_days(info.get(INFO_ROW_WORKLOAD, "")),
        "workload_note": info.get(INFO_ROW_WORKLOAD, NOT_SPECIFIED) or NOT_SPECIFIED,
        "must_deliver": must_deliver_items,
        "definition_of_done": dod_items,
        "quantitative_thresholds": builder.thresholds,
        "waves": {
            key: {
                "label": waves[key]["label"],
                "date_range": waves[key]["date_range"],
                "team_goal": waves[key]["team_goal"],
                "completion_condition": waves[key]["completion_condition"],
                "days": wave_days[key],
            }
            for key in ("A", "B", "C")
        },
        "code_freeze": code_freeze,
        "handoff": [text for text, _index in handoff_items],
        "conflict_boundary": " ".join(conflict_parts),
        "source": source,
        "requirements": builder.requirements,
    }
    return document


# ---------------------------------------------------------------------------
# 文档级元数据
# ---------------------------------------------------------------------------


def detect_document_version(content: DocxContent) -> str:
    """
    从文档中识别版本号（如 V3.0）。

    只接受文档自己写出的版本字符串；识别失败返回 not_specified。
    """
    for rows in content.tables:
        for row in rows:
            for cell in row:
                match = re.search(r"\bV(\d+\.\d+)\b", cell)
                if match:
                    return "V" + match.group(1)
    for _index, _style, text in content.paragraphs:
        match = re.search(r"\bV(\d+\.\d+)\b", text)
        if match:
            return "V" + match.group(1)
    return NOT_SPECIFIED


def detect_execution_period(content: DocxContent) -> str:
    """
    从文档中识别执行周期（如 2026-07-27 至 2026-08-09）。

    优先使用文档中出现的完整区间表述；识别失败返回 not_specified。
    """
    pattern = re.compile(r"(20\d{2})\s*[-/年]\s*(\d{1,2})\s*[-/月]\s*(\d{1,2})")
    for _index, _style, text in content.paragraphs:
        if "至" in text or "-" in text:
            found = pattern.findall(text)
            if len(found) >= 2:
                first, last = found[0], found[-1]
                return "%s-%02d-%02d 至 %s-%02d-%02d" % (
                    first[0], int(first[1]), int(first[2]),
                    last[0], int(last[1]), int(last[2]),
                )
    for rows in content.tables:
        for row in rows:
            for cell in row:
                found = pattern.findall(cell)
                if len(found) >= 2:
                    first, last = found[0], found[-1]
                    return "%s-%02d-%02d 至 %s-%02d-%02d" % (
                        first[0], int(first[1]), int(first[2]),
                        last[0], int(last[1]), int(last[2]),
                    )
    return NOT_SPECIFIED


def count_by_category(requirements: Sequence[Dict[str, Any]]) -> Dict[str, int]:
    """统计一个任务内各类别要求数量，用于 source-manifest.json。"""
    counts: Dict[str, int] = {}
    for requirement in requirements:
        counts[requirement["category"]] = counts.get(requirement["category"], 0) + 1
    return dict(sorted(counts.items()))


def count_by_wave(requirements: Sequence[Dict[str, Any]]) -> Dict[str, int]:
    """统计一个任务内各 Wave 的要求数量。"""
    counts: Dict[str, int] = {}
    for requirement in requirements:
        counts[requirement["wave"]] = counts.get(requirement["wave"], 0) + 1
    return dict(sorted(counts.items()))


def dump_yaml(document: Dict[str, Any]) -> str:
    """
    以确定性方式序列化 YAML。

    sort_keys=False 保留字段语义顺序；allow_unicode=True 保留中文原文；
    width 放大避免自动折行造成的无意义 diff。
    """
    return yaml.safe_dump(
        document,
        allow_unicode=True,
        sort_keys=False,
        default_flow_style=False,
        width=4096,
    )


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------


def build_all(docx_path: Path) -> Tuple[Dict[str, Dict[str, Any]], Dict[str, Any]]:
    """
    解析 DOCX 并构建九个任务文档与源清单骨架。

    返回：
        (documents, source_meta)，documents 为 {task_id: yaml_dict}。
    """
    content = read_docx(docx_path)
    waves = parse_wave_milestones(content)
    owner_map = parse_owner_map(content)
    must_tables = find_must_deliver_tables(content)
    handoff_items = build_handoff_requirements(content)
    source_meta = {
        "docx_name": docx_path.name,
        "docx_sha256": sha256_file(docx_path),
        "docx_size_bytes": docx_path.stat().st_size,
        "document_version": detect_document_version(content),
        "execution_period": detect_execution_period(content),
    }
    documents = {
        task_id: extract_task(
            content=content,
            task_id=task_id,
            must_deliver_table_index=must_tables[task_id],
            waves=waves,
            owner_map=owner_map,
            handoff_items=handoff_items,
            source_meta=source_meta,
        )
        for task_id in TASK_IDS
    }
    return documents, source_meta


def build_manifest(
    documents: Dict[str, Dict[str, Any]],
    source_meta: Dict[str, Any],
    rendered: Dict[str, str],
) -> Dict[str, Any]:
    """
    构建 source-manifest.json。

    generated_at_utc 是唯一非确定性字段，校验器会忽略它；
    files[].sha256 覆盖九个 YAML 的实际内容，用于检测人工改动。
    """
    files = []
    for task_id in TASK_IDS:
        text = rendered[task_id]
        requirements = documents[task_id]["requirements"]
        files.append(
            {
                "path": "docs/governance/task-requirements/%s.yaml" % task_id,
                "task_id": task_id,
                "sha256": sha256_text(text),
                "size_bytes": len(text.encode("utf-8")),
                "requirements_total": len(requirements),
                "by_category": count_by_category(requirements),
                "by_wave": count_by_wave(requirements),
                "must_deliver_items": len(documents[task_id]["must_deliver"]),
                "definition_of_done_items": len(documents[task_id]["definition_of_done"]),
                "quantitative_thresholds": len(documents[task_id]["quantitative_thresholds"]),
                "handoff_items": len(documents[task_id]["handoff"]),
            }
        )
    return {
        "schema_version": 1,
        "repository": REPOSITORY,
        "source": {
            "docx_name": source_meta["docx_name"],
            "docx_sha256": source_meta["docx_sha256"],
            "docx_size_bytes": source_meta["docx_size_bytes"],
            "document_version": source_meta["document_version"],
            "execution_period": source_meta["execution_period"],
            "note": "源 DOCX 只在本机读取，不得提交进仓库。",
        },
        "generated_by": "scripts/captain/extract_task_requirements.py",
        "generated_at_utc": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "determinism_note": "重复执行时 files[].sha256 必须保持不变；generated_at_utc 不参与校验。",
        "files": files,
        "totals": {
            "tasks": len(TASK_IDS),
            "requirements": sum(item["requirements_total"] for item in files),
            "quantitative_thresholds": sum(item["quantitative_thresholds"] for item in files),
        },
    }


def write_outputs(
    out_dir: Path,
    rendered: Dict[str, str],
    manifest: Dict[str, Any],
    force: bool,
) -> List[str]:
    """
    写出九个 YAML 与 source-manifest.json。

    当既有文件内容与本次生成不一致且未传入 --force 时，拒绝覆盖并报错，
    避免静默丢弃人工修订（治理指令：不覆盖人工修改而不提示）。
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    conflicts: List[str] = []
    for task_id, text in rendered.items():
        target = out_dir / ("%s.yaml" % task_id)
        if target.exists() and target.read_text(encoding="utf-8") != text and not force:
            conflicts.append(str(target))
    if conflicts:
        raise SystemExit(
            "以下文件内容与本次提取不一致，如确认覆盖请加 --force：\n  " + "\n  ".join(conflicts)
        )
    written: List[str] = []
    for task_id, text in rendered.items():
        target = out_dir / ("%s.yaml" % task_id)
        target.write_text(text, encoding="utf-8")
        written.append(str(target))
    manifest_path = out_dir / "source-manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    written.append(str(manifest_path))
    return written


def run_check(out_dir: Path, rendered: Dict[str, str]) -> int:
    """
    --check 模式：只比较既有 YAML 与本次提取结果，不写盘。

    返回：
        0 表示完全一致（幂等），1 表示存在差异或缺失文件。
    """
    differences: List[str] = []
    for task_id, text in rendered.items():
        target = out_dir / ("%s.yaml" % task_id)
        if not target.exists():
            differences.append("missing: %s" % target)
        elif target.read_text(encoding="utf-8") != text:
            differences.append("changed: %s" % target)
    if differences:
        for line in differences:
            print(line)
        return 1
    print("check ok: %d task files identical" % len(rendered))
    return 0


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    """定义命令行接口；DOCX 路径必须显式传入，不使用任何内置默认路径。"""
    parser = argparse.ArgumentParser(
        description="Extract T01-T09 task requirements from the SAGE125 V3.0 teammate manual."
    )
    parser.add_argument("--docx", required=True, help="Path to the V3.0 .docx manual (read-only).")
    parser.add_argument(
        "--out-dir",
        default="docs/governance/task-requirements",
        help="Output directory for T0X.yaml and source-manifest.json.",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="Compare against existing files without writing (idempotency check).",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing files even if their content differs.",
    )
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    """
    命令行入口。

    退出码：
        0 成功；1 --check 检出差异；2 源文档缺失或结构异常。
    """
    args = parse_args(argv)
    docx_path = Path(args.docx).expanduser()
    out_dir = Path(args.out_dir)
    try:
        documents, source_meta = build_all(docx_path)
    except (FileNotFoundError, LookupError, ValueError) as exc:
        print("EXTRACTION FAILED: %s" % exc, file=sys.stderr)
        return 2
    rendered = {task_id: dump_yaml(documents[task_id]) for task_id in TASK_IDS}
    if args.check:
        return run_check(out_dir, rendered)
    manifest = build_manifest(documents, source_meta, rendered)
    written = write_outputs(out_dir, rendered, manifest, args.force)
    print("source docx      : %s" % source_meta["docx_name"])
    print("source sha256    : %s" % source_meta["docx_sha256"])
    print("document version : %s" % source_meta["document_version"])
    print("execution period : %s" % source_meta["execution_period"])
    for task_id in TASK_IDS:
        requirements = documents[task_id]["requirements"]
        by_wave = count_by_wave(requirements)
        print(
            "%s: total=%d A=%d B=%d C=%d FINAL=%d FREEZE=%d dod=%d metrics=%d"
            % (
                task_id,
                len(requirements),
                by_wave.get("A", 0),
                by_wave.get("B", 0),
                by_wave.get("C", 0),
                by_wave.get("FINAL", 0),
                by_wave.get("FREEZE", 0),
                len(documents[task_id]["definition_of_done"]),
                len(documents[task_id]["quantitative_thresholds"]),
            )
        )
    print("wrote %d files" % len(written))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
